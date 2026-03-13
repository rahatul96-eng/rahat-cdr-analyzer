import streamlit as st
import pandas as pd
from docx import Document
import io
import re

st.set_page_config(page_title="Investigator Pro", layout="wide")
st.title("🕵️ CDR & FB-Name Finder")

# ফাইল আপলোড (Word, Excel, Txt)
uploaded_file = st.file_uploader("Upload CDR File", type=["docx", "txt", "xlsx"])

def parse_txt(text):
    # Regex to find: MSISDN, B Party, Location, Date_time, IMEI
    pattern = r"MSISDN:\s*(\d+).*?B Party:\s*(\w+).*?location:\s*(.*?)\s*Date_time:\s*(\d{14}).*?(IMEI:.*?)\["
    matches = re.findall(pattern, text, re.DOTALL)
    return [{'Date_Time': m[3], 'A_Party': m[0], 'B_Party_Loc': f"{m[1]} / {m[2]}", 'Rest': m[4].strip()} for m in matches]

if uploaded_file:
    try:
        if uploaded_file.name.endswith('.xlsx'):
            df = pd.read_excel(uploaded_file)
        elif uploaded_file.name.endswith('.docx'):
            doc = Document(uploaded_file)
            txt = "\n".join([p.text for p in doc.paragraphs])
            df = pd.DataFrame(parse_txt(txt))
        else:
            txt = uploaded_file.read().decode('utf-8')
            df = pd.DataFrame(parse_txt(txt))

        if not df.empty:
            st.success("✅ Data Loaded!")

            # --- Target OSINT (FB & Name Finder) ---
            st.subheader("🔍 Find Name & FB ID")
            nums = df['A_Party'].unique() if 'A_Party' in df.columns else df.iloc[:,0].unique()
            target = st.selectbox("Select Number", nums)
            
            c1, c2, c3 = st.columns(3)
            with c1: st.link_button("🔵 Facebook Search", f"https://www.facebook.com/search/top/?q={target}")
            with c2: st.link_button("📞 Truecaller (Name)", f"https://www.truecaller.com/search/bd/{target}")
            with c3: st.link_button("🌐 Google Search", f"https://www.google.com/search?q={target}")

            # --- Table Display ---
            st.dataframe(df, use_container_width=True)

            # --- Word Report ---
            if st.button("📥 Download Grouped Word Report"):
                out_doc = Document()
                for a_num, group in df.groupby(df.columns[1] if len(df.columns)>1 else df.columns[0]):
                    out_doc.add_heading(f'A Party: {a_num}', level=1)
                    table = out_doc.add_table(rows=1, cols=len(df.columns))
                    table.style = 'Table Grid'
                    for i, col in enumerate(df.columns): table.rows[0].cells[i].text = str(col)
                    for _, row in group.iterrows():
                        r_cells = table.add_row().cells
                        for i, v in enumerate(row): r_cells[i].text = str(v)
                bio = io.BytesIO()
                out_doc.save(bio)
                st.download_button("Download File", bio.getvalue(), "Investigator_Report.docx")
    except Exception as e:
        st.error(f"Error: {e}")
